import shlex
import shutil
import subprocess
import tempfile
import re
from html import unescape
from pathlib import Path
from typing import Any, Optional

from app.config import get_settings


class DocumentParser:
    def parse(self, file_path: str, content: Optional[bytes] = None) -> dict[str, Any]:
        path = Path(file_path)
        provider = get_settings().document_parser_provider.lower()
        if provider in {"auto", "mineru"}:
            mineru_result = self._try_mineru(path)
            if mineru_result:
                return mineru_result

        text = ""
        if path.suffix.lower() in {".txt", ".csv", ".md"}:
            text = self._decode_text(content if content is not None else path.read_bytes())
        elif path.suffix.lower() in {".html", ".htm"}:
            text = self._try_html(path, content)
        elif path.suffix.lower() in {".docx"}:
            text = self._try_docx(path)
        elif path.suffix.lower() in {".pdf"}:
            text = self._try_pdf(path)

        return {
            "provider": "local-document-parser",
            "text": text,
            "markdown": text,
            "tables": [],
            "needs_review": path.suffix.lower() in {".pdf", ".doc", ".docx"} and not bool(text.strip()),
        }

    def _try_mineru(self, path: Path) -> Optional[dict[str, Any]]:
        if path.suffix.lower() not in {".pdf", ".doc", ".docx"}:
            return None
        markdown_candidates = self._mineru_markdown_candidates(path)
        for candidate in markdown_candidates:
            if candidate.exists():
                text = candidate.read_text(encoding="utf-8", errors="ignore")
                return {
                    "provider": "mineru-markdown",
                    "text": text,
                    "markdown": text,
                    "tables": [],
                    "needs_review": False,
                }
        generated = self._run_mineru_cli(path)
        if generated:
            return generated
        return None

    def _mineru_markdown_candidates(self, path: Path) -> list[Path]:
        return [
            path.with_suffix(".md"),
            path.parent / f"{path.stem}.md",
            path.parent / path.stem / "auto" / f"{path.stem}.md",
        ]

    def _run_mineru_cli(self, path: Path) -> Optional[dict[str, Any]]:
        settings = get_settings()
        command = settings.mineru_command.strip()
        with tempfile.TemporaryDirectory(prefix="mineru-") as output_dir:
            output_path = Path(output_dir)
            commands = self._mineru_commands(command, path, output_path)
            for args in commands:
                try:
                    completed = subprocess.run(
                        args,
                        check=False,
                        capture_output=True,
                        text=True,
                        timeout=settings.mineru_timeout_seconds,
                    )
                except Exception:
                    continue
                if completed.returncode != 0:
                    continue
                markdown_files = sorted(output_path.rglob("*.md"))
                if not markdown_files:
                    markdown_files = [candidate for candidate in self._mineru_markdown_candidates(path) if candidate.exists()]
                if markdown_files:
                    text = "\n\n".join(item.read_text(encoding="utf-8", errors="ignore") for item in markdown_files)
                    return {
                        "provider": "mineru-cli",
                        "text": text,
                        "markdown": text,
                        "tables": [],
                        "needs_review": False,
                    }
        return None

    def _mineru_commands(self, configured_command: str, path: Path, output_path: Path) -> list[list[str]]:
        if configured_command:
            return [
                [
                    part.format(input=str(path), output=str(output_path))
                    for part in shlex.split(configured_command)
                ]
            ]
        commands: list[list[str]] = []
        magic_pdf = shutil.which("magic-pdf")
        if magic_pdf:
            commands.append([magic_pdf, "-p", str(path), "-o", str(output_path), "-m", "auto"])
        mineru = shutil.which("mineru")
        if mineru:
            commands.append([mineru, "-p", str(path), "-o", str(output_path)])
        return commands

    def _try_docx(self, path: Path) -> str:
        try:
            from docx import Document  # type: ignore
        except Exception:
            return ""
        try:
            document = Document(str(path))
            paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
            table_rows = []
            for table in document.tables:
                for row in table.rows:
                    table_rows.append(" | ".join(cell.text.strip() for cell in row.cells))
            return "\n".join([*paragraphs, *table_rows])
        except Exception:
            return ""

    def _try_pdf(self, path: Path) -> str:
        candidates = [self._try_pdf_pypdf(path), self._try_pdf_pymupdf(path)]
        return max(candidates, key=len, default="")

    def _try_pdf_pypdf(self, path: Path) -> str:
        try:
            from pypdf import PdfReader  # type: ignore
        except Exception:
            return ""
        try:
            reader = PdfReader(str(path))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception:
            return ""

    def _try_pdf_pymupdf(self, path: Path) -> str:
        try:
            import fitz  # type: ignore
        except Exception:
            return ""
        try:
            document = fitz.open(str(path))
            return "\n".join(page.get_text() or "" for page in document)
        except Exception:
            return ""

    def _try_html(self, path: Path, content: Optional[bytes] = None) -> str:
        raw = self._decode_text(content if content is not None else path.read_bytes())
        raw = re.sub(r"(?is)<(script|style).*?>.*?</\1>", "\n", raw)
        raw = re.sub(r"(?i)<br\s*/?>|</p>|</div>|</tr>|</h[1-6]>", "\n", raw)
        text = re.sub(r"(?s)<[^>]+>", "", raw)
        lines = []
        for line in unescape(text).splitlines():
            cleaned = " ".join(line.split())
            if cleaned:
                lines.append(cleaned)
        return "\n".join(lines)

    def _decode_text(self, content: bytes) -> str:
        for encoding in ("utf-8-sig", "utf-8", "gb18030"):
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        return ""


def get_document_parser() -> DocumentParser:
    return DocumentParser()
